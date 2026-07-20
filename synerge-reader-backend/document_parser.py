import io
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

MAX_FILE_BYTES = 50 * 1024 * 1024
MAX_PDF_PAGES = 150
MAX_EXTRACTED_CHARS = 300_000


class ExtractionError(Exception):
    def __init__(self, user_message: str, http_status: int = 422):
        self.user_message = user_message
        self.http_status = http_status
        super().__init__(user_message)


class UnsupportedFileTypeError(ExtractionError):
    def __init__(self, detail: str = ""):
        msg = "Unsupported file type. Please upload a PDF, DOCX, or plain text file."
        if detail:
            msg += f" ({detail})"
        super().__init__(user_message=msg, http_status=415)


@dataclass(frozen=True)
class ParsedPage:
    page_number: Optional[int]   # PDF: true 1-based source page number. DOCX/text: None.
    text: str


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]
    document_type: str          # "pdf", "docx", "text"
    page_count: int = 0
    char_count: int = 0
    truncated: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def file_type(self) -> str:
        return self.document_type


# Compatibility alias for existing type imports and result attribute
# access. ParsedDocument preserves the attributes used by current
# consumers.
ExtractionResult = ParsedDocument


def looks_like_text(content: bytes) -> bool:
    sample = content[:4096]
    if b"\x00" in sample:
        return False
    try:
        sample.decode("utf-8")
        return True
    except UnicodeDecodeError:
        pass
    if not sample:
        return False
    printable = sum(
        1 for b in sample
        if b == 0x09 or b == 0x0A or b == 0x0D or 0x20 <= b <= 0x7E
    )
    return (printable / len(sample)) > 0.85


def _detect_file_type(filename: str, content: bytes) -> str:
    if content[:5] == b"%PDF-":
        return "pdf"
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            names = set(z.namelist())
            if "[Content_Types].xml" in names and "word/document.xml" in names:
                return "docx"
    except zipfile.BadZipFile:
        pass
    if Path(filename).suffix.lower() in {".txt", ".md", ".csv"}:
        return "text"
    if looks_like_text(content):
        return "text"
    return "unknown"


def _check_zip_safety(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            infos = z.infolist()
            if len(infos) > 1000:
                raise ExtractionError("File is too large or complex to process.", 422)
            total = sum(info.file_size for info in infos)
            if total > 200 * 1024 * 1024:
                raise ExtractionError("File is too large or complex to process.", 422)
    except ExtractionError:
        raise
    except Exception:
        pass


def sanitize_filename(filename: "str | None") -> str:
    if not filename:
        return "untitled"
    name = Path(filename).name
    name = "".join(c for c in name if c.isprintable() and c not in r'/\\<>:"|?*')
    return name[:255] or "untitled"


def _truncate_pages(
    pages: list[ParsedPage],
    limit: int,
) -> tuple[list[ParsedPage], bool]:
    """Truncate PDF pages so that '\\n\\n'.join(text) stays within limit.

    PDF-only: keeps true source page numbers intact for every page that
    survives (partially or fully) the cut. DOCX/text use their own
    single-string truncation to stay character-for-character identical to
    the pre-refactor output.
    """
    sep = "\n\n"
    combined_len = 0
    new_pages = []
    for page in pages:
        prefix_len = len(sep) if new_pages else 0
        piece_len = prefix_len + len(page.text)
        if combined_len + piece_len <= limit:
            new_pages.append(page)
            combined_len += piece_len
        else:
            remaining = limit - combined_len - prefix_len
            if remaining > 0:
                new_pages.append(ParsedPage(page_number=page.page_number, text=page.text[:remaining]))
            return new_pages, True
    return new_pages, False


def _extract_pdf(content: bytes, filename: str) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as e:
        raise ExtractionError("PDF processing is not available on the server.", 500) from e

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        pdf_pages = pdf.pages
        if len(pdf_pages) > MAX_PDF_PAGES:
            raise ExtractionError(
                f"PDF exceeds the {MAX_PDF_PAGES}-page limit. Please upload a shorter document.",
                422,
            )

        pages = []
        for i, page in enumerate(pdf_pages, start=1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(ParsedPage(page_number=i, text=page_text))

        if not pages:
            raise ExtractionError(
                "This PDF appears to be scanned or image-based. Text extraction is not supported for image-only PDFs.",
                422,
            )

        pages, truncated = _truncate_pages(pages, MAX_EXTRACTED_CHARS)
        text = "\n\n".join(p.text for p in pages)

        return ParsedDocument(
            pages=pages,
            document_type="pdf",
            page_count=len(pdf_pages),
            char_count=len(text),
            truncated=truncated,
            warnings=["Document truncated to 300,000 characters."] if truncated else [],
        )


def _extract_docx(content: bytes, filename: str) -> ParsedDocument:
    try:
        import docx as python_docx
    except ImportError as e:
        raise ExtractionError("DOCX processing is not available on the server.", 500) from e

    _check_zip_safety(content)

    document = python_docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    truncated = False
    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]
        truncated = True

    if not text.strip():
        raise ExtractionError("This Word document appears to be empty or contains only images.", 422)

    pages = [ParsedPage(page_number=None, text=text)]

    return ParsedDocument(
        pages=pages,
        document_type="docx",
        char_count=len(text),
        truncated=truncated,
        warnings=["Document truncated to 300,000 characters."] if truncated else [],
    )


def _extract_text(content: bytes, filename: str) -> ParsedDocument:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="ignore")

    truncated = False
    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]
        truncated = True

    if not text.strip():
        raise ExtractionError("Uploaded text file is empty.", 422)

    pages = [ParsedPage(page_number=None, text=text)]

    return ParsedDocument(
        pages=pages,
        document_type="text",
        char_count=len(text),
        truncated=truncated,
        warnings=["Document truncated to 300,000 characters."] if truncated else [],
    )


def extract_text_from_upload(filename: str, content: bytes) -> ParsedDocument:
    if len(content) == 0:
        raise ExtractionError("Uploaded file is empty.", 422)
    if len(content) > MAX_FILE_BYTES:
        raise ExtractionError("File exceeds the 50 MB size limit.", 413)

    file_type = _detect_file_type(filename, content)

    if file_type == "unknown":
        raise UnsupportedFileTypeError()

    try:
        if file_type == "pdf":
            return _extract_pdf(content, filename)
        elif file_type == "docx":
            return _extract_docx(content, filename)
        else:
            return _extract_text(content, filename)
    except ExtractionError:
        raise
    except Exception:
        raise ExtractionError(
            "Failed to extract text from this file. The file may be corrupted.", 422
        )
